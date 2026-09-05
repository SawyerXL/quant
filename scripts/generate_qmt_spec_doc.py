"""
QMT主策略说明 → Word 生成器。
用法: python scripts/generate_qmt_spec_doc.py
输出: docs/QMT主策略说明.docx (源: docs/qmt_strategy_spec.md, 唯一权威口径)
支持: #/##/### 标题、>引用、markdown表格、```代码块、-列表、数字列表、**加粗**、`行内代码`
"""
import re
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

SRC = Path("docs/qmt_strategy_spec.md")
OUT = Path("docs/QMT主策略说明.docx")

BODY_FONT = "微软雅黑"
MONO_FONT = "Consolas"
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MID_BLUE = RGBColor(0x2F, 0x54, 0x96)
GRAY = RGBColor(0x59, 0x59, 0x59)

INLINE_RE = re.compile(r"(\*\*.*?\*\*|`[^`]*`)")

doc = Document()

# A4 页面
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.2)
sec.left_margin = sec.right_margin = Cm(2.2)

# 标题样式基线
for name, size, color, bold in [("Heading 1", 16, DARK_BLUE, True),
                                ("Heading 2", 13, MID_BLUE, True),
                                ("Heading 3", 11, MID_BLUE, True)]:
    st = doc.styles[name]
    st.font.name = BODY_FONT
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = color
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def set_run(run, bold=False, mono=False, size=None, color=None, italic=False):
    run.font.name = MONO_FONT if mono else BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_rich(p, text, size=None, color=None, base_bold=False):
    """按行内markdown(**加粗**/`代码`)拆run写入段落。"""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            r = p.add_run(part[2:-2])
            set_run(r, bold=True, size=size, color=color)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_run(r, mono=True, size=(size or 9) - 0.5, color=color)
        else:
            r = p.add_run(part)
            set_run(r, bold=base_bold, size=size, color=color)


def shade(cell_or_par, fill):
    """单元格或段落背景色。"""
    tcPr = cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, "_tc") \
        else cell_or_par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_md_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return
    header, body = rows[0], rows[1:]
    ncol = len(header)
    tbl = doc.add_table(rows=1 + len(body), cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rich(p, h, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(cell, "2F5496")
    for r_idx, row in enumerate(body):
        for c_idx, val in enumerate(row[:ncol]):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich(p, val, size=8)
            if r_idx % 2 == 1:
                shade(cell, "F2F5FA")
    doc.add_paragraph()  # 表后空行


def add_code_block(code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade(p, "F2F2F2")
    for i, ln in enumerate(code_lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(ln if ln.strip() else " ")
        set_run(r, mono=True, size=8)
    doc.add_paragraph()


# ── 逐行解析 ──
lines = SRC.read_text(encoding="utf-8").splitlines()
i = 0
while i < len(lines):
    ln = lines[i]
    s = ln.strip()
    if s.startswith("# ") and not s.startswith("## "):
        doc.add_heading(s[2:], level=1)
    elif s.startswith("## "):
        doc.add_heading(s[3:], level=2)
    elif s.startswith("### "):
        doc.add_heading(s[4:], level=3)
    elif s.startswith(">"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        add_rich(p, s.lstrip("> "), size=9, color=GRAY)
    elif s.startswith("```"):
        j = i + 1
        block = []
        while j < len(lines) and not lines[j].strip().startswith("```"):
            block.append(lines[j])
            j += 1
        add_code_block(block)
        i = j
    elif s.startswith("|"):
        j = i
        block = []
        while j < len(lines) and lines[j].strip().startswith("|"):
            block.append(lines[j])
            j += 1
        add_md_table(block)
        i = j - 1
    elif re.match(r"^\d+\.\s", s):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        add_rich(p, s, size=10)
    elif s.startswith(("- ", "✅ ", "⚠️ ", "⏸ ")):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.6)
        add_rich(p, s.lstrip("- "), size=10)
    elif not s:
        pass
    i += 1

# 生成信息页脚
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = meta.add_run(f"生成时间: {datetime.now():%Y-%m-%d %H:%M} | 源: {SRC}")
set_run(r, size=7.5, color=GRAY)

doc.save(OUT)
print(f"已生成: {OUT} ({OUT.stat().st_size/1024:.0f} KB)")
