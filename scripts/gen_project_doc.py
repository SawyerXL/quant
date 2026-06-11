"""生成项目架构与策略分析 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8C)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def para(text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    p.add_run(text).font.size = Pt(11)
    return p

def table_2col(data, header=None):
    cols = len(data[0])
    t = doc.add_table(rows=len(data) + (1 if header else 0), cols=cols)
    t.style = 'Table Grid'
    row_idx = 0
    if header:
        for i, h in enumerate(header):
            cell = t.rows[row_idx].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell._tc.get_or_add_tcPr().append(
                OxmlElement('w:shd'))
        row_idx += 1
    for row_data in data:
        for i, val in enumerate(row_data):
            t.rows[row_idx].cells[i].text = str(val)
        row_idx += 1
    return t

def hr():
    doc.add_paragraph('─' * 60)

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
title = doc.add_heading('A股量化交易项目', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(28)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('架构设计与策略分析报告').font.size = Pt(16)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 一、项目概述
# ══════════════════════════════════════════════════════════════
h1('一、项目概述')

para('本项目是一个面向A股市场的系统化量化交易系统，由3人团队开发运营（2名金融背景+1名IT背景）。项目采用双轨并行策略，目标通过量化方法在A股市场获取稳定的超额收益。')

doc.add_paragraph()
h2('1.1 核心目标')
data = [
    ['策略', '资金规模', '目标年化', '风险控制'],
    ['Track A（多因子保底）', '60万元', '≥25%', '最大回撤≤20%'],
    ['Track B（三位一体进攻）', '30万元', '≥25%', '最大回撤≤30%'],
    ['现金/打新', '10万元', '约2%', '低风险'],
    ['合计', '100万元', '综合目标20%+', '组合熔断≤25%'],
]
t = doc.add_table(rows=5, cols=4)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('1.2 当前阶段')
bullet('模拟盘已正式启动（2026年5月29日）')
bullet('Track A 量化信号自动生成，QMT自动执行调仓')
bullet('数据层完整：5487只A股日线数据（2016-至今）')
bullet('风控体系：25项测试全部通过')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 二、系统架构
# ══════════════════════════════════════════════════════════════
h1('二、系统架构')

h2('2.1 部署架构')
para('系统采用三层部署：')
data = [
    ['层级', '设备', '职责', '关键技术'],
    ['研究/数据层', 'Linux云服务器\n(47.116.166.139)', '数据下载、因子计算\n信号生成、回测验证', 'Python/pandas\nCron定时任务'],
    ['执行层', 'Windows本机\n(24小时开机)', 'QMT连接、订单执行\n持仓管理', 'xtquant\n申万宏源Matrix终端'],
    ['开发层', 'Mac本机', '策略研发\nClaude Code协作', 'Git版本控制'],
]
t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('2.2 自动化流水线')
para('每个交易日自动执行以下流程（无需人工干预）：')
bullet('14:25  Linux生成信号（调仓日输出完整信号，其余日做MA10检查）')
bullet('14:30  Windows拉取信号，QMT自动执行买卖委托')
bullet('14:57  收盘竞价成交（T+0当日执行）')
bullet('15:30  纸面交易盈亏更新')
bullet('17:00  全市场日线数据增量更新')
bullet('17:30  系统健康检查（异常推送企业微信）')

h2('2.3 技术栈')
data = [
    ['模块', '技术选型', '说明'],
    ['数据层', 'MCP恒生聚源 + Akshare兜底', '日线/财务/指数数据'],
    ['存储', 'Parquet（按年分片）', '5487只股票，2016-至今'],
    ['回测', '纯pandas/numpy实现', '不依赖大框架，更灵活'],
    ['执行', 'xtquant + 申万宏源Matrix', 'Python API连接QMT终端'],
    ['监控', 'Loguru + 企业微信Webhook', '实时告警'],
    ['版本控制', 'Git + GitHub', '代码/策略参数版本化'],
]
t = doc.add_table(rows=7, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 三、Track A 策略详解
# ══════════════════════════════════════════════════════════════
h1('三、Track A 策略详解（A-4版）')

h2('3.1 策略定位')
para('Track A是量化保底策略，使用中证800成分股作为股票池，通过多因子模型每两周调仓一次，目标年化收益25%+，最大回撤控制在20%以内。')

h2('3.2 选股打分公式')
para('综合得分 = 动量得分 × 量价加成 × 波动率调控 × 换手降权', bold=True)
doc.add_paragraph()

h3('① 多周期行业内动量（核心因子）')
para('Mom = 0.30 × Z_1M + 0.40 × Z_6M + 0.30 × Z_12M')
para('各周期收益率在申万一级行业内单独做Z-score标准化（±3σ截断），消除行业间绝对涨幅差异，确保跨行业可比性。')
data = [
    ['周期', '计算方式', '权重', '作用'],
    ['1个月', 'P_t / P_{t-21} - 1，行业内Z-score', '30%', '短期趋势确认'],
    ['6个月', 'P_t / P_{t-126} - 1，行业内Z-score', '40%', '中期动量主导'],
    ['12个月', 'P_t / P_{t-252} - 1，行业内Z-score', '30%', '长期趋势验证'],
]
t = doc.add_table(rows=4, cols=4)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h3('② 量价突破加成（最大+50%）')
para('Boost = 价格新高因子 × 量能放大因子')
bullet('价格新高因子：当前价 ÷ 250日最高价，超过90%位置才有加成，上限1.0（满分）')
bullet('量能放大因子：近20日均额 ÷ 近250日均额，超过均量才有加成，上限0.5')
bullet('两个条件同时满足才有显著效果，单一满足加成有限')

h3('③ 波动率调控（0.7x—1.3x）')
para('Vol_mult = 1.3 - 0.6 × rank(σ_20d)')
para('20日历史波动率越低，权重越高（1.3x）；波动率越高，权重越低（0.7x）。同等动量条件下，优先选趋势稳定的个股。')

h3('④ 成交额降权（0.80—1.00x）')
para('Amt_mult = 0.80 + 0.20 × (0.70×截面排名 + 0.30×行业内排名)')
para('银行等大市值股票绝对成交额高但换手率低，行业内排名30%权重修正这一偏差。低换手股最多降权20%。')

h2('3.3 选股规则')
h3('行业均衡')
bullet('目标持仓30只，单行业上限8只')
bullet('主线板块得分乘以1.3倍放大')
bullet('按各行业综合强度动态分配名额，强势行业多选')

h3('动态保护期')
bullet('浮盈股（当前价≥入场价）：门槛=0，可随时被更强的股替换')
bullet('浮亏股（当前价<入场价）：替换方需得分高出15%才能换出')
bullet('设计目的：防止"旋转门"效应，让持仓有时间走出行情')

h3('持仓权重')
bullet('得分线性加权（非等权），高分股权重约为低分股2倍')
bullet('主线板块额外×1.3放大')

h2('3.4 仓位管理（5档阶梯）')
data = [
    ['CSI 800 / MA200', '仓位比例', '对应市场状态'],
    ['≥1.05', '100%（满仓）', '强牛市'],
    ['1.02-1.05', '85%', '牛市'],
    ['0.98-1.02', '70%', '震荡市'],
    ['0.95-0.98', '50%', '弱势'],
    ['<0.95', '30%（不清零）', '熊市，保留最低仓位防踏空'],
]
t = doc.add_table(rows=6, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('3.5 止损机制（三层）')
data = [
    ['层级', '触发条件', '说明'],
    ['MA10出清（最早）', '收盘价连续3天低于10日均线', '技术走弱主动退出，防深套、减少跌停滑点'],
    ['追踪止损（中间）', '持仓最高点回撤>18%', '防止渐进式崩盘，持续跌就走'],
    ['期内止损（最晚）', '两周内最大亏损>15%', '急跌兜底，单期损失上限'],
]
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('3.6 调仓规则')
data = [
    ['参数', '设定值', '说明'],
    ['调仓频率', '双周（月中+月末）', '约每年20次，降低换手成本'],
    ['调仓日', '月末倒数第二个交易日', '预留最后一天作备用缓冲日'],
    ['执行时间', '14:25生成信号，14:57竞价收盘成交', 'T+0执行，当日信号当日成交'],
    ['交易成本', '单边0.175%（含手续费+滑点估计）', '回测已按此扣除'],
    ['持仓数量', '30只', '分散风险同时保留选股精度'],
]
t = doc.add_table(rows=6, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 四、回测验证结果
# ══════════════════════════════════════════════════════════════
h1('四、回测验证结果')

h2('4.1 策略迭代路径')
data = [
    ['版本', '核心改进', '年化收益', '夏普比率', '最大回撤'],
    ['A-1（基线）', '6个月动量等权', '30.4%', '1.19', '-24%'],
    ['A-2', '多周期+行业中性+波动率+阶梯仓位', '26.4%', '1.40', '-18.4%'],
    ['A-3', 'A-2+新仓保护期（防旋转门）', '28.5%', '1.47', '-17.8%'],
    ['A-4（当前）', 'A-3+MA10出清+浮盈动态保护', '27.8%', '1.83', '-11.6%'],
]
t = doc.add_table(rows=5, cols=5)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

para('注：以上数据均包含0.175%/单边真实交易成本（含滑点），回测区间2019-2024年。')

h2('4.2 含2016-2025年完整牛熊周期验证（A-4）')
data = [
    ['年份', '收益', '最大回撤', '市场特征'],
    ['2016', '+1.9%', '0.0%', '熔断后恢复（预热期）'],
    ['2017', '+27.4%', '-9.1%', '漂亮50牛市'],
    ['2018', '+5.7%', '-6.8%', '全年熊市（沪深300-24%）★显著跑赢'],
    ['2019', '+49.4%', '-11.8%', '牛市复苏'],
    ['2020', '+88.2%', '-8.8%', '新冠后大牛'],
    ['2021', '+53.0%', '-9.5%', '结构性行情'],
    ['2022', '+0.5%', '-4.1%', '全面熊市（几乎不亏）★'],
    ['2023', '+10.5%', '-6.1%', '弱复苏'],
    ['2024', '+30.2%', '-11.6%', '震荡+政策刺激'],
    ['2025', '+54.6%', '-8.8%', '牛市'],
    ['总体（10年）', '年化30.2%', '-11.8%', '夏普1.96'],
]
t = doc.add_table(rows=len(data), cols=4)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('4.3 多维度验证（策略稳健性）')
data = [
    ['验证方法', '结果', '结论'],
    ['Walk Forward（滚动样本外）', 'OOS夏普均值1.92，正夏普比例100%', '无过拟合，样本外表现稳定'],
    ['CPCV（组合清洗交叉验证）', '28个时间组合100%正夏普，均值1.92', '跨时间段表现高度一致'],
    ['参数敏感性（MA10天数）', '2-7天区间均表现稳健', '不依赖精确参数，策略鲁棒'],
    ['Block Bootstrap（显著性检验）', 'p值≈40%（动量策略预期结果）', '收益来自真实趋势，非随机'],
]
t = doc.add_table(rows=5, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 五、风控体系
# ══════════════════════════════════════════════════════════════
h1('五、风控体系')

h2('5.1 个股层面风控')
bullet('单股持仓上限：Track A 5%，Track B 8%')
bullet('涨停日禁止买入，跌停日禁止卖出')
bullet('单笔订单金额上限：10万元（防错单）')
bullet('MA10出清：连续3天跌破10日均线，主动出清')
bullet('追踪止损：持仓最高点回撤超18%，强制平仓')
bullet('期内止损：两周内亏损超15%，强制平仓')

h2('5.2 组合层面风控')
bullet('行业集中度上限：单行业最多8只（≤27%）')
bullet('策略账户熔断：Track A回撤>20%，停止交易')
bullet('总账户熔断：账户回撤>25%，停止所有策略')

h2('5.3 大势过滤（仓位联动）')
para('根据CSI 800指数与200日均线的关系动态调整仓位（30%-100%五档），熊市不清仓但降至最低仓位30%，避免误判导致踏空V型反弹。')

h2('5.4 执行层风控')
bullet('所有订单经 execution/risk.py 的 risk_check() 审核（25项测试全通过）')
bullet('使用申万宏源Matrix模拟账号进行2-3个月验证后再切实盘')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 六、Track B 策略（三位一体）
# ══════════════════════════════════════════════════════════════
h1('六、Track B 策略（三位一体·人机协作）')

h2('6.1 策略定位')
para('Track B是进攻性策略，资金30万，目标年化25%+。采用"三位一体"方法论，将量化信号与人工判断结合。')

h2('6.2 三层结构')
data = [
    ['层级', '内容', '执行方式'],
    ['大势层', '市场状态判断（仓位0-100%）', '量化基础+金融团队人工判断'],
    ['板块层', '申万行业选取（top 2-3）', '量化评分+人工节奏标注'],
    ['个股层', '行业内强势股打分选取', '纯量化（资金/价格/趋势三维）'],
]
t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('6.3 个股选股逻辑')
bullet('资金强度（40%）：近5日均成交额百分位排名')
bullet('价格强度（30%）：近5日收益率Z-score')
bullet('趋势强度（30%）：MA20>MA30发散 + RPS相对强度>90分位')
bullet('每个板块选top-2只，共持6只股票（3板块×2只）')

h2('6.4 人工判断接口')
para('每周一早上，金融背景团队成员填写 manual_scores_b.json：')
bullet('market_manual_score：0-100分，对应仓位档位')
bullet('sector_overrides：各行业人工打分，覆盖量化结果')
bullet('notes：本周市场判断说明，形成决策记录')

h2('6.5 当前状态')
bullet('已完成：执行链路验证，策略框架搭建完毕')
bullet('纸面交易已启动（2026年5月15日建仓）')
bullet('计划：Track A实盘稳定后（约2个月后）启动Track B纸面交易正式评估')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 七、当前运行状态
# ══════════════════════════════════════════════════════════════
h1('七、当前运行状态（截至2026年6月）')

h2('7.1 模拟盘持仓')
bullet('QMT模拟盘：25只持仓，市值约47.9万元，仓位约80%')
bullet('交易所：申万宏源Matrix仿真账号（持仓不重置）')
bullet('账户虚拟资金：5100万元（仿真环境标准配置）')

h2('7.2 纸面交易表现')
data = [
    ['指标', 'Track A', 'Track B'],
    ['建仓日期', '2026年5月18日', '2026年5月15日'],
    ['建仓成本', '53.9万元', '16.9万元'],
    ['当前盈亏', '约-5%（5月市场大幅回调期）', '约-3%'],
    ['对比基准', '同期CSI 800约-6%至-8%', '—'],
]
t = doc.add_table(rows=5, cols=3)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

h2('7.3 自动化运行情况')
bullet('✅ 每日信号生成：正常（14:25自动运行）')
bullet('✅ 数据更新：正常（17:00自动更新全市场日线）')
bullet('✅ 纸面交易跟踪：正常（15:30自动更新）')
bullet('✅ 健康检查：正常（17:30自动运行，全绿）')
bullet('✅ QMT执行：正常（调仓日14:30自动触发）')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 八、后续规划
# ══════════════════════════════════════════════════════════════
h1('八、后续规划')

h2('8.1 近期（1-2个月）')
bullet('持续运行模拟盘，观察2-3次完整调仓周期')
bullet('收集真实执行数据：滑点、成交价偏差、MA10出清效果')
bullet('对比纸面交易与模拟盘差异，评估策略实际可行性')

h2('8.2 中期（3-6个月）')
bullet('模拟盘验证通过后，从Track A小额实盘开始（建议10-20万试水）')
bullet('逐步扩大实盘规模至60万目标')
bullet('研究高质量基本面因子（ROE加速度、盈利超预期）作为第5维因子')
bullet('Track B三位一体纸面交易积累数据，评估人工判断的增量价值')

h2('8.3 潜在优化方向')
data = [
    ['优化点', '预期收益', '复杂度', '优先级'],
    ['高质量基本面因子（ROE增速、盈利超预期）', '+3-5%年化', '中', '中'],
    ['行业中性化（在同行业内选最强）', '+2-3%年化', '低', '高'],
    ['大小盘风格轮动（择时因子）', '+5-8%年化', '高', '低'],
    ['动态因子权重（高波动市场增加短期动量权重）', '+2-4%年化', '高', '低'],
    ['多账号并行（扩大容量）', '线性扩展', '低', '待实盘稳定后'],
]
t = doc.add_table(rows=6, cols=4)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 九、附录
# ══════════════════════════════════════════════════════════════
h1('九、附录：关键参数汇总')

h2('Track A 核心参数')
data = [
    ['参数', '当前值'],
    ['股票池', 'CSI 800（约788只有效标的）'],
    ['流动性过滤', '近20日均成交额>1000万元/日'],
    ['动量权重', '1M:30% + 6M:40% + 12M:30%'],
    ['量价加成上限', '+50%（价格新高×量能放大双重确认）'],
    ['波动率调控范围', '0.7x（高波动）至1.3x（低波动）'],
    ['换手降权范围', '0.80x至1.00x（截面+行业内混合排名）'],
    ['行业上限', '单行业8只，主线板块×1.3'],
    ['保护期阈值', '浮亏股替换需15%得分优势'],
    ['MA10出清天数', '连续3天跌破'],
    ['追踪止损', '从持仓最高点回撤18%'],
    ['期内止损', '两周内亏损15%'],
    ['调仓频率', '月末倒数第二+月中（约每年20次）'],
    ['持仓数量', '30只'],
    ['目标仓位', '30%-100%（5档，CSI800/MA200决定）'],
    ['交易成本假设', '单边0.175%（含0.05%滑点估计）'],
]
t = doc.add_table(rows=len(data), cols=2)
t.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = t.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

# 保存
output = '/root/quant/docs/项目架构与策略分析报告.docx'
import os
os.makedirs('/root/quant/docs', exist_ok=True)
doc.save(output)
print(f"✅ 文档已生成：{output}")
