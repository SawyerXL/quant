"""
生成 QMT 主策略说明 Word 文档（v2.2）。
内容来源: docs/qmt_strategy_spec.md（唯一权威口径）。
用法: python scripts/generate_qmt_spec_docx.py
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("docs/QMT主策略说明_v2.2.docx")


def set_cn_font(run, size=10.5, bold=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), "微软雅黑")


def h(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_cn_font(run, size={1: 16, 2: 13, 3: 11.5}[level], bold=True)
    return p


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, bold=bold)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_cn_font(run)
    return p


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htext in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htext)
        set_cn_font(run, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_cn_font(run)
    return t


def main():
    doc = Document()
    title = doc.add_heading("", level=0)
    tr = title.add_run("QMT 主策略完整说明")
    set_cn_font(tr, size=22, bold=True)
    st = para(doc, "版本 v2.2 · 2026-08-30 · 唯一权威口径（与 docs/qmt_strategy_spec.md 同步）")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "所有参数改动必须回测 A/B 对比 + 用户确认（CLAUDE.md 工程纪律）。")

    h(doc, "1. 一句话定位")
    para(doc, "成交额 TOP60 等权动量池 + MA200 择时控仓 + MA10 退出纪律 + 拥挤度过滤。"
              "熊市赚“少亏+反弹回血”，牛市赚“动量+弹性”。没有做空，没有绝对收益机制。")

    h(doc, "2. 核心机制")
    h(doc, "2.1 股票池：成交额 TOP60（v2.2 定案，收益优先）", 2)
    bullet(doc, "取最近 20 个交易日平均成交额前 60 名，等权")
    bullet(doc, "2026-08-30 池子规模 A/B 定案：TOP60+过滤全期年化 +6.92%（vs TOP30+过滤 +6.15%），夏普 0.49 最高；TOP80 全面平庸")
    bullet(doc, "优点：熊市反弹日弹性大（8 月池子等权 +12.9% vs 上证 +1.6%）；代价：动量拥挤（8/19 型跌停潮）")

    h(doc, "2.2 择时控仓：CSI800 / MA200 五档", 2)
    table(doc,
          ["指数/MA200", "仓位", "状态"],
          [["≥1.05", "100%", "强牛"], ["1.02~1.05", "85%", "牛"],
           ["0.98~1.02", "70%", "震荡"], ["0.95~0.98", "50%", "弱"],
           ["<0.95", "30%", "熊（不清零，防踏空V反）"]])
    para(doc, "8 月实测：平均仓位 47%，8/19 跌停潮只伤净值 -1.4%（而非满仓 -3.1%）。")

    h(doc, "2.3 退出纪律：MA10-4d", 2)
    para(doc, "单票连续 4 个交易日收盘在 MA10 下方 → 清仓；双周调仓自动补买。单票层面截断持续走弱。")

    h(doc, "2.4 止盈：TP30/60（诚实标注：装饰性）", 2)
    para(doc, "浮盈 +30% 卖 1/3、+60% 再卖 1/3。2026-08-29 复验：现口径下从未改变任何一笔交易"
              "（票在 +30% 前就被调仓/MA10 换掉），与“无止盈”回测逐位相同。保留为状态机字段。")

    h(doc, "2.5 拥挤度过滤（v2.1 新增）", 2)
    para(doc, "调仓时剔除 20 日波动率 >5% 的票（只用 T-1 及以前数据）。"
              "动机：8/19 失血——9 只跌停凶手事前波动率 4.4~7.9% vs 防御票 1.0~1.7%；"
              "凶手 20 日涨幅仅 +2~15%，涨幅过滤拦不住。")
    table(doc,
          ["窗口", "基线", "+过滤", "差"],
          [["2019-2022", "+8.86%", "+11.23%", "+2.4pp"],
           ["全期 2019-2026.8", "+5.10%", "+6.92%", "+1.8pp"],
           ["近段 2022-2026.8", "+3.87%", "+5.24%", "+1.4pp"],
           ["近段 2024.7-2026.8", "+11.30%", "+16.65%", "+5.4pp"]])
    para(doc, "稳健性：阈值网格 4~6% 平滑平台；严格口径（不含当日）与含当天差 +0.19pp。")

    h(doc, "2.6 止损：成本 -15% 全卖（V2）", 2)
    para(doc, "Windows 端 stop_monitor 执行（20 分钟心跳 --go）。追踪止损已停用（只砍赢家、对回撤零贡献）。")

    h(doc, "3. 完整参数表")
    table(doc,
          ["参数", "值", "说明"],
          [["pool_size", "60（v2.2 定案）", "实盘端待同步（仍跑 TOP30）"],
           ["rebalance_freq", "双周（15日+月末）", "周频更差、月频回撤低但收益差"],
           ["ma_exit_days", "4", "MA10 连续跌破"],
           ["take_profit", "30%/60% 各卖 1/3", "装饰性"],
           ["absolute/trailing stop", "关（实盘 -15%）", "7/8 消融：对池子有害"],
           ["过热过滤（涨幅类）", "关", "7/8 消融：纯破坏"],
           ["max_vol20", "5.0（v2.1）", "拥挤度过滤"],
           ["vol20_use_today", "False", "严格口径"],
           ["max_position_pct", "10%", "新票仓位上限"],
           ["commission", "0.13% 双边", "佣金+印花+滑点"],
           ["cash_yield", "2%/年", "低仓位现金收益"],
           ["min_bars", "250", "新股过滤"]])

    h(doc, "4. 执行链（数据流）")
    bullet(doc, "Windows QMT（仅执行）：15:10 每日持仓快照（修复目标，现状飘 09:15~16:58）；调仓指令 buy/sell 增量（非幂等）；stop_monitor 20 分钟心跳")
    bullet(doc, "Linux 服务器：qmt_nav_track.py —— 100 万 notional 口径；快照价仅采信 ≥15:00，否则本地收盘价重估 → logs/qmt_nav_history.parquet")

    h(doc, "5. 风控红线（CLAUDE.md 强制）")
    bullet(doc, "单票仓位上限（Track A 口径）、单笔订单 5 万、涨停禁买/跌停禁卖")
    bullet(doc, "账户回撤熔断 25%；所有订单过 execution/risk.py 的 risk_check()")

    h(doc, "6. 历史验证")
    para(doc, "回测（TOP60 口径，2019-2026.8）：全期年化 +5.1%（+过滤 +6.9%），夏普 0.31（+过滤 0.49），回撤 -27.1%。"
              "年度：2020 +26.6%、2025 +40.3%、2022/2023 亏损（熊市防守年）。")
    para(doc, "8 月实盘（100 万 notional，收盘价重估口径）：7/20→8/28 +3.78%，夏普 2.11，回撤 -4.26%，平均仓位 47%。"
              "同期上证 -0.09%、中证800 -3.54%。峰值 8/17 +7.05 万 → 收尾 +3.78 万，失血主因 8/19 单日 -3.1%（9 只持仓跌停）。")
    para(doc, "诚实对照：8/3 持仓满仓死拿 +12.9%，策略只兑现 +6.6%——纪律成本 6.3pp，"
              "买的是“躲开下一次 8/19 是趋势起点”。")

    h(doc, "7. 已知问题 / 待办")
    bullet(doc, "实盘端同步：Windows 仍跑 TOP30 无过滤——待隧道恢复后改为 TOP60 + 波动率过滤")
    bullet(doc, "Windows 快照时点飘移：Linux 侧已有重估防御，根治需登 Windows 改计划任务")
    bullet(doc, "做T增厚：t0 开关保持关——验证期（9/23 满月）目前 53.8% 胜率/-0.99% 不达标")
    bullet(doc, "止盈开关与“无止盈”回测逐位相同——装饰性，不构成收益来源")

    h(doc, "8. 变更日志")
    table(doc,
          ["日期", "变更", "依据"],
          [["2026-07-08", "消融：砍入场/过热/追踪/绝对止损，仅 MA10-4d + TP30/60", "年化 5.7%→9.3%"],
           ["2026-07-09", "追踪止损停用，改成本 -15%", "追踪只砍赢家"],
           ["2026-07-20", "TP v3 固定分批 30/60", "A/B（8/29 复验未复现）"],
           ["2026-08-29", "净值口径：快照价仅采信 ≥15:00，否则收盘价重估", "用户质疑 +6.71% 虚高"],
           ["2026-08-30", "拥挤度过滤 max_vol20=5%", "四窗口 A/B + 网格 + 严格口径复验"],
           ["2026-08-30", "池子规模定案 TOP60（v2.2，收益优先）", "池子 A/B：60 收益/夏普最高"]])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    main()
